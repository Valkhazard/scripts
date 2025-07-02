import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from PIL import Image
from docx2pdf import convert
import subprocess
import math


RUTA_PRINCIPAL = r"C:\Users\user\Documents\Cuenta de cobro Amorchis"

CARPETA_WORD = os.path.join(RUTA_PRINCIPAL, 'WORD')
CARPETA_PDF = os.path.join(RUTA_PRINCIPAL, 'PDF')

EXT_IMAGEN = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp'}

os.makedirs(CARPETA_WORD, exist_ok=True)
os.makedirs(CARPETA_PDF, exist_ok=True)

# configuracion de margenes y tamanos
# pagina A4: 21cm x 29.7cm
# margenes tipicos de Word: 2.5cm por lado
PAGINA_ANCHO_CM = 21.0
PAGINA_ALTO_CM = 29.7
MARGEN_HORIZONTAL_CM = 2.5
MARGEN_VERTICAL_CM = 2.5
ESPACIO_ENTRE_IMAGENES_CM = 0.5

# area util para imagenes
AREA_UTIL_ANCHO_CM = PAGINA_ANCHO_CM - (2 * MARGEN_HORIZONTAL_CM)  # ~16cm
AREA_UTIL_ALTO_CM = PAGINA_ALTO_CM - (2 * MARGEN_VERTICAL_CM) - 4  # ~23cm (restamos espacio para titulo)

# configuracion para diferentes layouts
LAYOUT_1_IMG = {
    'max_width': AREA_UTIL_ANCHO_CM * 0.9,  # 90% del ancho util
    'max_height': AREA_UTIL_ALTO_CM * 0.8   # 80% del alto util
}

LAYOUT_2_IMG = {
    'max_width': AREA_UTIL_ANCHO_CM * 0.48,  # 48% del ancho util cada una
    'max_height': AREA_UTIL_ALTO_CM * 0.8
}

LAYOUT_4_IMG = {
    'max_width': AREA_UTIL_ANCHO_CM * 0.48,  # 48% del ancho util
    'max_height': (AREA_UTIL_ALTO_CM * 0.4) - ESPACIO_ENTRE_IMAGENES_CM  # 40% del alto util menos espacio
}

def get_image_dimensions_cm(img_path):
    """obtiene las dimensiones reales de la imagen en cm"""
    try:
        with Image.open(img_path) as img:
            width_px, height_px = img.size
            # obtener DPI, usar 96 como default si no esta disponible
            dpi = img.info.get('dpi', (96, 96))
            if isinstance(dpi, tuple):
                dpi = dpi[0]
            
            # convertir pixeles a cm
            width_cm = width_px / dpi * 2.54
            height_cm = height_px / dpi * 2.54
            
            return width_cm, height_cm, width_px / height_px  # aspect ratio
    except Exception as e:
        print(f"Error al leer dimensiones de {img_path}: {e}")
        return 10.0, 10.0, 1.0  # valores por defecto

def calcular_dimensiones_optimas(img_path, max_width_cm, max_height_cm, min_width_cm=4.0):
    """calcula las dimensiones optimas para una imagen respetando limites y aspect ratio"""
    width_cm, height_cm, aspect_ratio = get_image_dimensions_cm(img_path)
    
    # calcular escalas necesarias
    escala_ancho = max_width_cm / width_cm if width_cm > max_width_cm else 1.0
    escala_alto = max_height_cm / height_cm if height_cm > max_height_cm else 1.0
    
    # usar la escala mas restrictiva para mantener proporciones
    escala = min(escala_ancho, escala_alto)
    
    # aplicar escala
    new_width = width_cm * escala
    new_height = height_cm * escala
    
    # verificar ancho minimo
    if new_width < min_width_cm:
        factor_min = min_width_cm / new_width
        new_width = min_width_cm
        new_height = new_height * factor_min
        
        # si despues del ajuste minimo excede el alto maximo, ajustar proporcionalmente
        if new_height > max_height_cm:
            factor_alto = max_height_cm / new_height
            new_height = max_height_cm
            new_width = new_width * factor_alto
    
    return new_width, new_height

def evaluar_capacidad_pagina(imagenes_grupo):
    """evalua si un grupo de imagenes puede caber en una pagina"""
    if len(imagenes_grupo) <= 1:
        return True
    
    # calcular dimensiones que tendria cada imagen en el layout de 4
    dimensiones_calculadas = []
    for img in imagenes_grupo:
        width_cm, height_cm = calcular_dimensiones_optimas(
            img, 
            LAYOUT_4_IMG['max_width'], 
            LAYOUT_4_IMG['max_height']
        )
        dimensiones_calculadas.append((width_cm, height_cm))
    
    # verificar si pueden organizarse en cuadricula 2x2
    if len(imagenes_grupo) <= 4:
        # calcular altura total necesaria para 2 filas
        fila1_altura = max(dimensiones_calculadas[0][1], dimensiones_calculadas[1][1] if len(dimensiones_calculadas) > 1 else 0)
        fila2_altura = 0
        if len(dimensiones_calculadas) > 2:
            fila2_altura = max(dimensiones_calculadas[2][1], dimensiones_calculadas[3][1] if len(dimensiones_calculadas) > 3 else 0)
        
        altura_total = fila1_altura + fila2_altura + ESPACIO_ENTRE_IMAGENES_CM
        
        # verificar si cabe en el area util
        return altura_total <= AREA_UTIL_ALTO_CM
    
    return False

def agrupar_imagenes_inteligente(imagenes):
    """agrupa imagenes de manera inteligente priorizando 4 por pagina cuando sea posible"""
    if len(imagenes) <= 1:
        return [imagenes]
    
    grupos = []
    imagenes_restantes = imagenes.copy()
    
    # procesar de 4 en 4 mientras sea posible
    while len(imagenes_restantes) >= 4:
        grupo_candidato = imagenes_restantes[:4]
        
        # verificar si estas 4 imagenes pueden caber juntas
        if evaluar_capacidad_pagina(grupo_candidato):
            grupos.append(grupo_candidato)
            imagenes_restantes = imagenes_restantes[4:]
        else:
            # si no caben 4, intentar con 2
            grupo_candidato = imagenes_restantes[:2]
            if evaluar_capacidad_pagina(grupo_candidato):
                grupos.append(grupo_candidato)
                imagenes_restantes = imagenes_restantes[2:]
            else:
                # si ni siquiera caben 2, poner 1 sola
                grupos.append([imagenes_restantes[0]])
                imagenes_restantes = imagenes_restantes[1:]
    
    # procesar imagenes restantes
    if len(imagenes_restantes) > 0:
        if len(imagenes_restantes) <= 2:
            # si quedan 1 o 2, verificar si pueden ir con el grupo anterior
            if grupos and len(grupos[-1]) + len(imagenes_restantes) <= 4:
                # verificar si el grupo combinado puede caber
                grupo_combinado = grupos[-1] + imagenes_restantes
                if evaluar_capacidad_pagina(grupo_combinado):
                    grupos[-1] = grupo_combinado
                else:
                    grupos.append(imagenes_restantes)
            else:
                grupos.append(imagenes_restantes)
        else:
            # si quedan 3, verificar si caben todas juntas
            if evaluar_capacidad_pagina(imagenes_restantes):
                grupos.append(imagenes_restantes)
            else:
                # dividir en grupos mas pequenos
                grupos.append(imagenes_restantes[:2])
                grupos.append(imagenes_restantes[2:])
    
    return grupos

def agregar_imagen_al_documento(doc, img_path, layout_config, centrada=True):
    """agrega una imagen al documento con la configuracion especificada"""
    try:
        width_cm, height_cm = calcular_dimensiones_optimas(
            img_path, 
            layout_config['max_width'], 
            layout_config['max_height']
        )
        
        p = doc.add_paragraph()
        if centrada:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm), height=Cm(height_cm))
        
        return True
    except Exception as e:
        print(f"Error al agregar imagen {img_path}: {e}")
        return False

def agregar_espacio_entre_imagenes(doc, espacio_cm=0.3):
    """agrega un pequeno espacio entre imagenes"""
    p = doc.add_paragraph()
    p.space_after = Pt(espacio_cm * 28.35)  # convertir cm a puntos

def verificar_imagenes_nuevas(doc, lista_imagenes):
    """verifica que imagenes no estan ya en el documento para evitar duplicados"""
    # obtener nombres de archivos de imagenes que queremos agregar
    nombres_imagenes_nuevas = [os.path.basename(img) for img in lista_imagenes]
    
    # por ahora, devolver todas las imagenes (se puede mejorar para detectar duplicados)
    # en una version mas avanzada, se podria analizar el contenido del documento
    return lista_imagenes

def extraer_titulo_comentario(doc):
    """
    extrae el titulo (primer heading o parrafo centrado grande) y el comentario (primer parrafo centrado despues del titulo)
    del documento Word. devuelve (titulo, comentario).
    """
    titulo = None
    comentario = None
    # buscar heading nivel 0 o 1 como titulo preferente
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if not titulo:
                titulo = para.text.strip()
            continue
        # si no hay heading, buscar centrados grandes
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER and para.text.strip():
            if not titulo:
                titulo = para.text.strip()
            elif not comentario:
                comentario = para.text.strip()
                break
    return titulo, comentario

def insertar_imagenes_dos_por_pagina(doc, imagenes):
    """
    inserta imagenes en el documento, 2 por pagina. si ambas imagenes son largas (alto/ancho > 1.2), las coloca en horizontal (lado a lado).
    si no, usa la logica vertical (una encima de otra). si hay una sola imagen, ocupa toda la pagina.
    siempre asegura que los pares de imagenes esten juntos en la misma pagina.
    """
    n = len(imagenes)
    if n == 0:
        return
    idx = 0
    while idx < n:
        grupo = imagenes[idx:idx+2]
        if len(grupo) == 1:
            # solo una imagen en la pagina, usar todo el alto util
            ancho_util = AREA_UTIL_ANCHO_CM
            alto_util = AREA_UTIL_ALTO_CM
            width_cm, height_cm, aspect = get_image_dimensions_cm(grupo[0])
            escala_ancho = ancho_util / width_cm
            escala_alto = alto_util / height_cm
            escala = min(escala_ancho, escala_alto, 1.0)
            new_width = width_cm * escala
            new_height = height_cm * escala
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(grupo[0], width=Cm(new_width), height=Cm(new_height))
        elif len(grupo) == 2:
            # detectar si ambas son largas
            w1, h1, _ = get_image_dimensions_cm(grupo[0])
            w2, h2, _ = get_image_dimensions_cm(grupo[1])
            es_larga_1 = h1 / w1 > 1.2
            es_larga_2 = h2 / w2 > 1.2
            if es_larga_1 and es_larga_2:
                # layout horizontal: lado a lado
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                cell1, cell2 = table.rows[0].cells
                ancho_img = AREA_UTIL_ANCHO_CM / 2
                alto_img = AREA_UTIL_ALTO_CM
                for img, cell in zip(grupo, [cell1, cell2]):
                    width_cm, height_cm, _ = get_image_dimensions_cm(img)
                    escala_ancho = ancho_img / width_cm
                    escala_alto = alto_img / height_cm
                    escala = min(escala_ancho, escala_alto, 1.0)
                    new_width = width_cm * escala
                    new_height = height_cm * escala
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img, width=Cm(new_width), height=Cm(new_height))
            else:
                # layout vertical: una encima de otra
                for img in grupo:
                    ancho_util = AREA_UTIL_ANCHO_CM
                    alto_util = AREA_UTIL_ALTO_CM / 2
                    width_cm, height_cm, _ = get_image_dimensions_cm(img)
                    escala_ancho = ancho_util / width_cm
                    escala_alto = alto_util / height_cm
                    escala = min(escala_ancho, escala_alto, 1.0)
                    new_width = width_cm * escala
                    new_height = height_cm * escala
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img, width=Cm(new_width), height=Cm(new_height))
        idx += 2
        # solo agregar salto de pagina si hay mas imagenes despues
        if idx < n:
            doc.add_page_break()

carpetas = [d for d in os.listdir(RUTA_PRINCIPAL)
            if os.path.isdir(os.path.join(RUTA_PRINCIPAL, d)) and d not in ['WORD', 'PDF']]

carpetas_no_procesadas = []
carpetas_sin_imagenes_para_agregar = []

for carpeta in carpetas:
    ruta_carpeta = os.path.join(RUTA_PRINCIPAL, carpeta)
    nombre_word = f"{carpeta}.docx"
    ruta_word = os.path.join(ruta_carpeta, nombre_word)  # SIEMPRE en la subcarpeta
    ruta_word_copia = os.path.join(CARPETA_WORD, nombre_word)  # copia para PDF

    # obtener lista de imagenes
    imagenes = []
    nombres_vistos = set()
    for archivo in os.listdir(ruta_carpeta):
        nombre, ext = os.path.splitext(archivo)
        if ext.lower() in EXT_IMAGEN and archivo not in nombres_vistos:
            imagenes.append(os.path.join(ruta_carpeta, archivo))
            nombres_vistos.add(archivo)
    imagenes.sort()

    # si ya existe el documento Word en la subcarpeta
    if os.path.exists(ruta_word):
        doc = Document(ruta_word)
        titulo_existente, comentario_existente = extraer_titulo_comentario(doc)
        print(f"\n[INFO] Documento existente: {nombre_word}")
        print(f"  Titulo detectado: {titulo_existente}")
        print(f"  Comentario detectado: {comentario_existente}")
        if imagenes:
            print(f"Anadiendo imagenes al documento existente: {carpeta}")
            imagenes_a_agregar = verificar_imagenes_nuevas(doc, imagenes)
            if not imagenes_a_agregar:
                print(f"No hay imagenes nuevas para agregar en: {carpeta}")
                doc.save(ruta_word)
                shutil.copy2(ruta_word, ruta_word_copia)
                continue
            print(f"Se agregaran {len(imagenes_a_agregar)} imagenes nuevas")
            insertar_imagenes_dos_por_pagina(doc, imagenes_a_agregar)
            doc.save(ruta_word)
            print(f"Imagenes anadidas exitosamente a: {ruta_word}")
        else:
            carpetas_sin_imagenes_para_agregar.append(carpeta)
            doc.save(ruta_word)
        shutil.copy2(ruta_word, ruta_word_copia)
        continue

    # si no hay imagenes, saltar
    if not imagenes:
        continue

    # crear nuevo documento en la subcarpeta
    print(f"\nProcesando carpeta: {carpeta}")
    titulo = input("Titulo para el documento: ")
    comentario = input("Comentario para el documento: ")

    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(MARGEN_VERTICAL_CM)
        section.bottom_margin = Cm(MARGEN_VERTICAL_CM)
        section.left_margin = Cm(MARGEN_HORIZONTAL_CM)
        section.right_margin = Cm(MARGEN_HORIZONTAL_CM)
    doc.add_heading(titulo, 0)
    if comentario.strip():
        p_comentario = doc.add_paragraph()
        run_comentario = p_comentario.add_run(comentario)
        run_comentario.font.name = 'Arial Black'
        run_comentario.font.size = Pt(16)
        p_comentario.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_comentario.paragraph_format.space_after = Pt(14)
    total_imgs = len(imagenes)
    print(f"Procesando {total_imgs} imagenes...")
    if total_imgs > 0:
        insertar_imagenes_dos_por_pagina(doc, imagenes)
    doc.save(ruta_word)
    print(f"Documento Word creado: {ruta_word}")
    shutil.copy2(ruta_word, ruta_word_copia)

# convertir Word a PDF
print("\nAbriendo documentos Word para revision...")
docs_word = [f for f in os.listdir(CARPETA_WORD) if f.lower().endswith('.docx')]

for docx_file in docs_word:
    ruta_docx = os.path.join(CARPETA_WORD, docx_file)
    try:
        subprocess.Popen(['start', '', ruta_docx], shell=True)
    except Exception as e:
        print(f"No se pudo abrir {ruta_docx}: {e}")

# borrar los documentos Word generados
while True:
    borrar_word = input("¿Quieres borrar los documentos Word generados antes de continuar? (s/N): ").strip().lower()
    if borrar_word == 's':
        for docx_file in docs_word:
            ruta_docx = os.path.join(CARPETA_WORD, docx_file)
            try:
                os.remove(ruta_docx)
                print(f"Borrado: {ruta_docx}")
            except Exception as e:
                print(f"No se pudo borrar {ruta_docx}: {e}")
        print("\nTodos los documentos Word han sido borrados. puedes volver a ejecutar el script si necesitas modificar algo.")
        exit(0)
    elif borrar_word == 'n' or borrar_word == '':
        print("No se borraron los documentos Word generados.")
        break
    else:
        print("Por favor, responde solo con 's' para borrar o 'n' para no borrar.")

input("\n¿Deseas continuar con la conversion a PDF? (Presiona Enter para continuar o Ctrl+C para cancelar)")

print("Convirtiendo documentos a PDF...")
for docx_file in docs_word:
    ruta_docx = os.path.join(CARPETA_WORD, docx_file)
    ruta_pdf = os.path.join(CARPETA_PDF, docx_file.replace('.docx', '.pdf'))
    try:
        convert(ruta_docx, ruta_pdf)
        print(f"Convertido: {docx_file} -> PDF")
    except Exception as e:
        print(f"Error convirtiendo {docx_file} a PDF: {e}")

# reportes finales
print("\n" + "="*50)
print("REPORTE FINAL")
print("="*50)

if carpetas_no_procesadas:
    print("\nCarpetas no procesadas (ya tenian Word):")
    for c in carpetas_no_procesadas:
        print(f"- {c}")

if carpetas_sin_imagenes_para_agregar:
    print("\nCarpetas con Word existente sin imagenes para agregar:")
    for c in carpetas_sin_imagenes_para_agregar:
        print(f"- {c}")

print(f"\nDocumentos Word procesados: {len(docs_word)}")
print(f"Ubicacion de archivos Word: {CARPETA_WORD}")
print(f"Ubicacion de archivos PDF: {CARPETA_PDF}")
print("\n¡Proceso completado!")